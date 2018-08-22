import docx
from docx.shared import Pt

# class to create the Word templates
class Template:

    # constructor takes a student object and file name as aruments.
    def __init__(self, user, file_name):
        self.user = user
        self.file_name = file_name
        self.document = docx.Document('../files/template.docx') # this is preset location and should not change
        # TODO: Create option to import template from file location.

    # this method creates the template in a Word document.
    def create_template(self):
        doc = self.document # set the document to a variable

        # Space variables used for formating the page correctly
        space1 = " " * 40
        space2 = " " * 76
        space3 = " " * 40
        space4 = space1 + (" " * 28) + space2
        space5 = " " * 165
        space6 = space5 + (" " * 50)

        # instantiate an Instructor object.
        instructor = Instructor()

        # Each doc.paragraphs operation below is removing the existing line at
        # the index and replacing it with proper spacing and text data from the
        # student object passed into the constructor.
        doc.paragraphs[5].clear()
        doc.paragraphs[5].add_run((" " * 5) + self.user.user)

        doc.paragraphs[6].clear()
        doc.paragraphs[6].add_run((" " * 5) + self.user.address1)

        doc.paragraphs[7].clear()
        doc.paragraphs[7].add_run((" " * 5) + self.user.address2)

        doc.paragraphs[8].clear()
        doc.paragraphs[8].add_run((" " * 5) + "{}, {} {}".format(self.user.city, self.user.state, self.user.zipcode))

        doc.paragraphs[11].clear()
        doc.paragraphs[11].add_run((" " * 5) + space1 + self.user.date + space2 + self.user.facility)

        doc.paragraphs[12].clear()
        doc.paragraphs[12].add_run((" " * 5) + space1 + self.user.exp_date + space2 + instructor.name)

        doc.paragraphs[13].clear()
        doc.paragraphs[13].add_run((" " * 5) + space4 + (" " * 4) + instructor.number)

        doc.paragraphs[15].clear()
        run = doc.paragraphs[15].add_run(space3 + self.user.user)
        font = run.font
        font.size = Pt(20)

        doc.paragraphs[26].clear()
        doc.paragraphs[26].add_run((" " * 135 )+ self.user.user)

        doc.paragraphs[28].clear()
        run = doc.paragraphs[28].add_run(space5 + self.user.facility)
        font = run.font
        font.size = Pt(9)

        doc.paragraphs[29].clear()
        run = doc.paragraphs[29].add_run(space5 + self.user.date)
        font = run.font
        font.size = Pt(9)

        doc.paragraphs[30].clear()
        run = doc.paragraphs[30].add_run(space5 + self.user.exp_date + (" " * 53 ) + str(self.user.course_hours))
        font = run.font
        font.size = Pt(9)

        doc.paragraphs[31].clear()
        doc.paragraphs[31].add_run((" " * 40) + instructor.number)

        # saves the document to a folder on the desktop.
        doc.save('../../../../cpr_templates/' + self.file_name + '.docx')
        # TODO: Create a method for saving the documents to a specified location

# used to print the Word template to the console for development purposes in
# getting the formatting of the generated template correct.
def print_template():
        doc = docx.Document('../files/template.docx')
        counter = 0
        for para in doc.paragraphs:
            print(para.text + str(counter))
            counter += 1
